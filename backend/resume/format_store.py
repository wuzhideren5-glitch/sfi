"""
Original File Store — saves uploaded resume files for format-preserving export.

When user uploads a PDF resume:
1. Original bytes saved to profiles/{user_id}_original.pdf
2. PDF metadata extracted: text blocks with position, font, size
3. AI edits are applied to the extracted structured data
4. Export: modify the original PDF directly (PyMuPDF text replacement)
"""
from __future__ import annotations

import io
import json
from pathlib import Path

import fitz  # PyMuPDF

PROFILES_DIR = Path(__file__).parent.parent / "profiles"
PROFILES_DIR.mkdir(exist_ok=True)


class OriginalStore:
    """Manages the original uploaded resume file and its format metadata."""

    def __init__(self, user_id: str):
        self.user_id = user_id
        self._pdf_path = PROFILES_DIR / f"{user_id}_original.pdf"
        self._meta_path = PROFILES_DIR / f"{user_id}_format_meta.json"

    # ═══════════════════════════════════════════════════════════
    # Save original
    # ═══════════════════════════════════════════════════════════

    def save_original(self, file_bytes: bytes, filename: str):
        """Save the original uploaded file and extract format metadata."""
        self._pdf_path.write_bytes(file_bytes)

        # Extract format metadata from PDF
        meta = self._extract_format_meta(file_bytes)
        self._meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2))

    def _extract_format_meta(self, file_bytes: bytes) -> list[dict]:
        """Extract text blocks with position and font info from PDF."""
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        blocks = []

        for page_num, page in enumerate(doc):
            # Get text blocks with detailed info
            text_blocks = page.get_text("dict")["blocks"]
            for block in text_blocks:
                if block.get("type") != 0:  # not text
                    continue
                for line in block.get("lines", []):
                    line_text = ""
                    spans_info = []
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                        spans_info.append({
                            "text": span.get("text", ""),
                            "font": span.get("font", "Helvetica"),
                            "size": round(span.get("size", 11), 1),
                            "bold": "Bold" in span.get("font", ""),
                            "color": self._rgb_to_hex(span.get("color", 0)),
                        })

                    if line_text.strip():
                        bbox = line["bbox"]
                        blocks.append({
                            "page": page_num,
                            "text": line_text.strip(),
                            "x": round(bbox[0], 1),
                            "y": round(bbox[1], 1),
                            "width": round(bbox[2] - bbox[0], 1),
                            "height": round(bbox[3] - bbox[1], 1),
                            "spans": spans_info,
                        })

        doc.close()
        return blocks

    @staticmethod
    def _rgb_to_hex(rgb_int: int) -> str:
        """Convert PyMuPDF color int to hex."""
        r = (rgb_int >> 16) & 0xFF
        g = (rgb_int >> 8) & 0xFF
        b = rgb_int & 0xFF
        return f"#{r:02x}{g:02x}{b:02x}" if rgb_int != 0 else "#000000"

    # ═══════════════════════════════════════════════════════════
    # Read
    # ═══════════════════════════════════════════════════════════

    def has_original(self) -> bool:
        return self._pdf_path.exists()

    def get_original_bytes(self) -> bytes | None:
        if self._pdf_path.exists():
            return self._pdf_path.read_bytes()
        return None

    def get_format_meta(self) -> list[dict]:
        if self._meta_path.exists():
            return json.loads(self._meta_path.read_text(encoding="utf-8"))
        return []

    # ═══════════════════════════════════════════════════════════
    # PDF Text Replacement (preserves exact formatting)
    # ═══════════════════════════════════════════════════════════

    def apply_text_replacements(self, replacements: list[dict]) -> bytes:
        """Apply text replacements to the original PDF, preserving formatting.

        Args:
            replacements: [{"old_text": "原文", "new_text": "新文本"}, ...]

        Returns:
            Modified PDF bytes with original formatting intact.
        """
        if not self._pdf_path.exists():
            return b""

        doc = fitz.open(str(self._pdf_path))

        for rep in replacements:
            old = rep.get("old_text", "")
            new = rep.get("new_text", "")
            if not old or old == new:
                continue

            for page in doc:
                # Find all instances of the old text
                areas = page.search_for(old)
                for rect in areas:
                    # Redact (remove) the old text
                    page.add_redact_annot(rect, fill=(1, 1, 1))  # white fill
                page.apply_redactions()

                # Insert new text at the position of the first match
                if areas:
                    rect = areas[0]
                    # Get font info from the area to match
                    font_size = abs(rect.y1 - rect.y0) * 0.75  # approximate
                    font_size = max(font_size, 9)
                    try:
                        page.insert_text(
                            (rect.x0, rect.y1 - 2),
                            new,
                            fontsize=font_size,
                            fontname="china-s",
                            color=(0, 0, 0),
                        )
                    except Exception:
                        # Fallback: use helv (ASCII only)
                        ascii_only = ''.join(c for c in new if ord(c) < 128)
                        if ascii_only:
                            page.insert_text(
                                (rect.x0, rect.y1 - 2),
                                ascii_only,
                                fontsize=font_size,
                            )

        buffer = io.BytesIO()
        doc.save(buffer, garbage=4, deflate=True)
        doc.close()
        buffer.seek(0)
        return buffer.getvalue()

    # ═══════════════════════════════════════════════════════════
    # DOCX from Original Format
    # ═══════════════════════════════════════════════════════════

    def export_docx_from_original(self, text_replacements: dict[str, str] | None = None) -> bytes:
        """Build DOCX mimicking the original PDF layout.

        Uses extracted format metadata to match font sizes, bold, and structure.
        text_replacements: {original_text: new_text} mapping for AI edits.
        """
        from docx import Document
        from docx.shared import Pt, Inches, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        meta = self.get_format_meta()
        if not meta:
            return b""

        doc = Document()

        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Track what we've added to avoid duplicates
        seen = set()
        replacements = text_replacements or {}

        for block in meta:
            text = block["text"]
            # Apply replacements
            for old_t, new_t in replacements.items():
                if old_t in text:
                    text = text.replace(old_t, new_t)

            if text in seen:
                continue
            seen.add(text)

            # Get font info from first span
            spans = block.get("spans", [{}])
            first_span = spans[0]
            font_size = first_span.get("size", 11)
            is_bold = first_span.get("bold", False)

            p = doc.add_paragraph()
            # Infer if this is a heading from font size and bold
            if font_size >= 16 and is_bold:
                # Name header
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(font_size)
            elif font_size >= 13 and is_bold:
                # Section heading
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(font_size)
                p.paragraph_format.space_before = Pt(8)
            elif is_bold:
                # Bold line (company name, role)
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(font_size)
            else:
                run = p.add_run(text)
                run.font.size = Pt(font_size)

            run.font.name = 'Calibri'

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
