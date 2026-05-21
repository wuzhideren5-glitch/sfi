"""
Resume Service — structured resume management with AI-driven editing.

Features:
- Parse uploaded PDF resume into structured MD format
- AI-driven natural language editing: "把这段改成更量化的表述"
- Format-preserving edits: AI modifies content, structure stays clean
- Export to DOCX (Word) and PDF
"""
from __future__ import annotations

import io
import json
import logging
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from core.llm import chat as deepseek_chat
from core.profile_store import ProfileStore
from core.session_store import TEST_USER_ID

logger = logging.getLogger(__name__)

EDIT_SYSTEM_PROMPT = """你是简历编辑助手。根据用户指令修改简历中指定部分的内容。

## 规则
1. **只修改被要求修改的内容**，其他部分保持原样
2. **保持格式结构**：如果原文是列表，返回列表；如果是段落，返回段落
3. **量化优先**：如果有数字、百分比、具体成果，用它们替换模糊表述
4. **STAR 法则**：Situation → Task → Action → Result
5. **输出纯文本**：只输出修改后的内容，不要 markdown 包裹，不加解释

## 示例
原文：在XX公司实习，负责数据分析工作
指令：更量化
输出：在XX公司担任数据分析实习生，处理10万+条用户数据，通过SQL优化将查询效率提升40%，输出3份业务分析报告"""


class ResumeService:
    """Manages resume storage and editing."""

    def __init__(self):
        self._store = ProfileStore(user_id=TEST_USER_ID)

    # ═══════════════════════════════════════════════════════════
    # Read / Write
    # ═══════════════════════════════════════════════════════════

    def get_resume_text(self) -> str:
        """Get the full resume section from MD profile."""
        return self._store.get_resume_section()

    def set_resume_text(self, content: str):
        """Set/replace the full resume section."""
        self._store.set_resume_section(content)

    def has_resume(self) -> bool:
        """Check if a resume exists."""
        content = self.get_resume_text()
        return bool(content and content.strip() and "<!--" not in content[:10])

    # ═══════════════════════════════════════════════════════════
    # Build from Parsed Profile
    # ═══════════════════════════════════════════════════════════

    def build_from_profile(self, profile: dict) -> str:
        """Build a formatted markdown resume from parsed profile data."""
        lines = []

        # Header
        name = profile.get("name", "")
        lines.append(f"# {name}")
        contact_parts = []
        if profile.get("city"):
            contact_parts.append(profile["city"])
        if profile.get("gender"):
            contact_parts.append(profile["gender"])
        if profile.get("age"):
            contact_parts.append(f"{profile['age']}岁")
        if contact_parts:
            lines.append(" | ".join(contact_parts))
        lines.append("")

        # Education
        edu = profile.get("education", [])
        if edu:
            lines.append("## 教育背景")
            for e in edu:
                parts = []
                school = e.get("school", "")
                major = e.get("major", "")
                degree = e.get("degree", "")
                year = e.get("year", "")
                if school:
                    parts.append(f"**{school}**")
                if major:
                    parts.append(major)
                if degree:
                    parts.append(degree)
                if year:
                    parts.append(year)
                lines.append("- " + " · ".join(parts))
            lines.append("")

        # Internships
        internships = profile.get("internships", [])
        if internships:
            lines.append("## 实习经历")
            for i in internships:
                company = i.get("company", "")
                role = i.get("role", "")
                duration = i.get("duration", "")
                desc = i.get("description", "")
                lines.append(f"### {role} | {company}")
                if duration:
                    lines.append(f"*{duration}*")
                if desc:
                    lines.append(f"- {desc}")
                lines.append("")
            lines.append("")

        # Projects
        projects = profile.get("projects", [])
        if projects:
            lines.append("## 项目经历")
            for p in projects:
                name = p.get("name", "")
                role = p.get("role", "")
                desc = p.get("description", "")
                lines.append(f"### {name}")
                if role:
                    lines.append(f"*{role}*")
                if desc:
                    lines.append(f"- {desc}")
                lines.append("")
            lines.append("")

        # Skills
        skills = profile.get("skills", [])
        if skills:
            lines.append("## 技能")
            lines.append("- " + "、".join(skills))
            lines.append("")

        # Certificates
        certs = profile.get("certificates", [])
        if certs:
            lines.append("## 证书与资质")
            lines.append("- " + "、".join(certs))
            lines.append("")

        content = "\n".join(lines)
        self._store.set_resume_section(content)
        return content

    # ═══════════════════════════════════════════════════════════
    # AI-Driven Editing
    # ═══════════════════════════════════════════════════════════

    async def edit_section(
        self,
        section: str,
        instruction: str,
        item_index: int | None = None,
    ) -> dict:
        """Edit a specific section of the resume using AI.

        Args:
            section: Which section to edit ("education", "internships", "projects", "skills", "summary")
            instruction: Natural language instruction
            item_index: Which item in the list to edit (if applicable)
        """
        full_resume = self.get_resume_text()
        if not full_resume:
            return {"success": False, "error": "简历为空，请先上传简历"}

        # Find the target section content
        section_content = self._extract_section_content(full_resume, section, item_index)
        if not section_content:
            return {"success": False, "error": f"未找到简历中的「{section}」部分"}

        before = section_content.strip()

        # Call LLM to edit
        try:
            edited = deepseek_chat(
                messages=[
                    {"role": "system", "content": EDIT_SYSTEM_PROMPT},
                    {"role": "user", "content": f"原文：\n{before}\n\n指令：{instruction}"},
                ],
                model="deepseek-chat",
                temperature=0.3,
                max_tokens=1000,
            )
            after = edited.strip()
        except Exception as e:
            logger.error("Resume edit failed: %s", e)
            return {"success": False, "error": f"AI编辑失败: {e}"}

        # Apply the edit back to the full resume
        new_resume = self._replace_section_content(full_resume, section, before, after, item_index)
        self.set_resume_text(new_resume)

        return {
            "success": True,
            "section": section,
            "before": before,
            "after": after,
            "changes": f"已修改「{section}」部分",
        }

    def _extract_section_content(
        self, full_resume: str, section: str, item_index: int | None = None
    ) -> str:
        """Extract a section's content from the resume markdown."""
        section_map = {
            "education": "教育背景",
            "internships": "实习经历",
            "projects": "项目经历",
            "skills": "技能",
            "certificates": "证书与资质",
            "summary": "个人总结",
        }
        heading = section_map.get(section, section)

        # Find the section
        pattern = rf'## {re.escape(heading)}\s*\n(.*?)(?=\n## |\Z)'
        m = re.search(pattern, full_resume, re.DOTALL)
        if not m:
            return ""

        content = m.group(1).strip()

        # If item_index is specified, extract only that item
        if item_index is not None and section in ("education", "internships", "projects"):
            items = self._split_list_items(content)
            if 0 <= item_index < len(items):
                return items[item_index]
            return ""

        return content

    def _replace_section_content(
        self,
        full_resume: str,
        section: str,
        old_content: str,
        new_content: str,
        item_index: int | None = None,
    ) -> str:
        """Replace section content in the resume, preserving formatting."""
        section_map = {
            "education": "教育背景",
            "internships": "实习经历",
            "projects": "项目经历",
            "skills": "技能",
            "certificates": "证书与资质",
            "summary": "个人总结",
        }
        heading = section_map.get(section, section)

        pattern = rf'## {re.escape(heading)}\s*\n(.*?)(?=\n## |\Z)'

        if item_index is not None and section in ("education", "internships", "projects"):
            # Replace single item within the section
            m = re.search(pattern, full_resume, re.DOTALL)
            if m:
                section_content = m.group(1)
                items = self._split_list_items(section_content)
                if 0 <= item_index < len(items):
                    items[item_index] = new_content
                    new_section = "\n\n".join(items)
                    return full_resume[:m.start(1)] + new_section + full_resume[m.end(1):]
        else:
            # Replace entire section
            replacement = f"## {heading}\n{new_content}"
            return re.sub(pattern, replacement, full_resume, flags=re.DOTALL)

        return full_resume

    def _split_list_items(self, content: str) -> list[str]:
        """Split section content into individual items (separated by ### headings or blank lines)."""
        # Split by ### subheadings
        items = re.split(r'\n(?=### )', content)
        if len(items) <= 1:
            # Try splitting by blank lines
            items = [i.strip() for i in content.split('\n\n') if i.strip()]
        return items

    # ═══════════════════════════════════════════════════════════
    # Export
    # ═══════════════════════════════════════════════════════════

    def export_docx(self) -> bytes:
        """Export resume as Word document."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        content = self.get_resume_text()
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            if line.startswith('# '):
                # Name / title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line[2:])
                run.bold = True
                run.font.size = Pt(18)

            elif line.startswith('## '):
                # Section heading
                p = doc.add_paragraph()
                run = p.add_run(line[3:])
                run.bold = True
                run.font.size = Pt(13)
                # Add a bottom border
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)

            elif line.startswith('### '):
                # Sub-heading (role, project name)
                p = doc.add_paragraph()
                run = p.add_run(line[4:])
                run.bold = True
                run.font.size = Pt(11)

            elif line.startswith('- '):
                # Bullet point
                p = doc.add_paragraph(line[2:], style='List Bullet')

            elif line.startswith('*') and line.endswith('*'):
                # Italic line (duration etc.)
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.italic = True
                run.font.size = Pt(10)

            else:
                # Normal text
                p = doc.add_paragraph(line)

            i += 1

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_pdf(self) -> bytes:
        """Export resume as PDF using markdown → HTML → PDF approach."""
        content = self.get_resume_text()

        # Simple HTML conversion
        html = self._md_to_html(content)

        # Use PyMuPDF or a simple approach
        try:
            # Try using weasyprint if available
            from weasyprint import HTML
            return HTML(string=html).write_pdf()
        except ImportError:
            pass

        # Fallback: use reportlab
        try:
            return self._export_pdf_reportlab(content)
        except Exception:
            pass

        # Last resort: return the markdown as text in a PDF
        return self._export_pdf_simple(content)

    def _md_to_html(self, md_text: str) -> str:
        """Convert markdown to basic HTML for PDF export."""
        lines = md_text.split('\n')
        html_parts = [
            '<!DOCTYPE html><html><head><meta charset="utf-8">',
            '<style>',
            'body { font-family: "PingFang SC", "Microsoft YaHei", sans-serif; font-size: 11pt; line-height: 1.6; max-width: 700px; margin: 40px auto; color: #333; }',
            'h1 { text-align: center; font-size: 18pt; margin-bottom: 4px; }',
            'h2 { font-size: 13pt; border-bottom: 2px solid #333; padding-bottom: 4px; margin-top: 20px; }',
            'h3 { font-size: 11pt; margin-bottom: 2px; }',
            'ul { margin: 4px 0; padding-left: 20px; }',
            'li { margin: 2px 0; }',
            'p { margin: 4px 0; }',
            'em { color: #666; }',
            '</style></head><body>',
        ]

        i = 0
        in_list = False
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                i += 1
                continue

            if line.startswith('# '):
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append(f'<h1>{self._escape_html(line[2:])}</h1>')
            elif line.startswith('## '):
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append(f'<h2>{self._escape_html(line[3:])}</h2>')
            elif line.startswith('### '):
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append(f'<h3>{self._escape_html(line[4:])}</h3>')
            elif line.startswith('- '):
                if not in_list:
                    html_parts.append('<ul>')
                    in_list = True
                html_parts.append(f'<li>{self._escape_html(line[2:])}</li>')
            elif line.startswith('*') and line.endswith('*'):
                html_parts.append(f'<p><em>{self._escape_html(line.strip("*"))}</em></p>')
            else:
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append(f'<p>{self._escape_html(line)}</p>')
            i += 1

        if in_list:
            html_parts.append('</ul>')
        html_parts.append('</body></html>')
        return '\n'.join(html_parts)

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def _export_pdf_reportlab(self, content: str) -> bytes:
        """Export using reportlab (fallback)."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm,
                                leftMargin=25*mm, rightMargin=25*mm)
        styles = getSampleStyleSheet()
        story = []

        for line in content.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 4*mm))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Title']))
            elif line.startswith('## '):
                story.append(Spacer(1, 6*mm))
                story.append(Paragraph(line[3:], styles['Heading2']))
                story.append(HRFlowable(width="100%", thickness=1, color="#333"))
            elif line.startswith('### '):
                story.append(Paragraph(f"<b>{line[4:]}</b>", styles['Normal']))
            elif line.startswith('- '):
                story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            elif line.startswith('*') and line.endswith('*'):
                story.append(Paragraph(f"<i>{line.strip('*')}</i>", styles['Normal']))
            else:
                story.append(Paragraph(line, styles['Normal']))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_pdf_simple(self, content: str) -> bytes:
        """Last resort: use fpdf2."""
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font('CJK', '', '/System/Library/Fonts/PingFang.ttc', uni=True)
            pdf.set_font('CJK', '', 11)

            for line in content.split('\n'):
                pdf.multi_cell(0, 6, line)

            return pdf.output(dest='S').encode('latin-1')
        except Exception:
            # Ultimate fallback: return PDF with just the text
            return content.encode('utf-8')

    # ═══════════════════════════════════════════════════════════
    # Chat-Driven Resume Editing
    # ═══════════════════════════════════════════════════════════

    async def auto_edit_from_chat(
        self,
        user_message: str,
        ai_reply: str,
    ) -> dict | None:
        """Analyze a conversation turn for resume-related requests and auto-edit.

        If the user asks to modify their resume, the AI extracts the edit
        instruction and applies it automatically.
        """
        # Check if the conversation is about resume editing
        triggers = ["简历", "resume", "cv", "经历", "修改", "改成", "调整为", "帮我写", "优化"]
        combined = f"{user_message} {ai_reply}"
        if not any(t in combined.lower() for t in triggers):
            return None

        # Ask LLM to extract edit instruction
        try:
            result = deepseek_chat(
                messages=[
                    {"role": "system", "content": """分析对话，判断是否涉及简历修改。如果是，提取修改指令。

返回 JSON：
{
  "is_edit": true/false,
  "section": "education/internships/projects/skills/summary",
  "instruction": "具体的修改指令",
  "item_index": null  // 如果是列表中的某一项，指定索引
}

如果不是简历修改请求，返回 {"is_edit": false}。只返回 JSON。"""},
                    {"role": "user", "content": f"学生：{user_message}\n小苗老师：{ai_reply}"},
                ],
                model="deepseek-chat",
                temperature=0.0,
                max_tokens=300,
            )

            result = result.strip()
            if result.startswith("```"):
                result = result.split("\n", 1)[-1].rsplit("```", 1)[0]

            data = json.loads(result)
            if not data.get("is_edit"):
                return None

            section = data.get("section", "")
            instruction = data.get("instruction", "")
            item_index = data.get("item_index")

            if not section or not instruction:
                return None

            edit_result = await self.edit_section(section, instruction, item_index)
            return edit_result

        except Exception as e:
            logger.debug("Auto-edit detection skipped: %s", e)
            return None
